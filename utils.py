import pandas as pd
from io import BytesIO
from docx import Document

def load_library_data(filename):
    """ Load library data from an Excel file """
    try:
        return pd.read_excel(filename)
    except Exception as e:
        print(f"Error loading {filename}: {e}")
        return None

def match_item_numbers(user_data):
    """ Identifies the lookup column in the user-uploaded file """
    possible_columns = ["Item No.", "Item Number", "Article No.", "Variant ID", "Description"]
    for sheet, df in user_data.items():
        df.columns = df.columns.str.strip().str.lower()
        for col in df.columns:
            if any(keyword.lower() in col for keyword in possible_columns):
                return col
    return None

def generate_product_list_presentation(user_data, lookup_col, library_df):
    """ Generate a Word document with matched products """
    document = Document()
    document.add_heading("Product List for Presentations", level=1)

    for sheet, df in user_data.items():
        df.columns = df.columns.str.strip().str.lower()
        quantity_col = next((col for col in df.columns if "quantity" in col.lower() or "qty" in col.lower()), None)
        
        if quantity_col:
            merged = df.merge(library_df, left_on=lookup_col, right_on="EUR Item No.", how="left")
            for _, row in merged.iterrows():
                if pd.notna(row["Product"]) and pd.notna(row[quantity_col]):
                    document.add_paragraph(f"{row[quantity_col]} X {row['Product']}")
    
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def generate_order_import(user_data, lookup_col):
    """ Generate an Excel file for order import """
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet, df in user_data.items():
            df.columns = df.columns.str.strip().str.lower()
            quantity_col = next((col for col in df.columns if "quantity" in col.lower() or "qty" in col.lower()), None)
            if quantity_col:
                df_filtered = df[[lookup_col, quantity_col]].dropna()
                df_filtered.to_excel(writer, index=False, header=False)
    return buffer.getvalue()

def generate_detailed_product_list(user_data, lookup_col, library_df):
    """ Generate an Excel file with detailed product information """
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet, df in user_data.items():
            df.columns = df.columns.str.strip().str.lower()
            quantity_col = next((col for col in df.columns if "quantity" in col.lower() or "qty" in col.lower()), None)
            merged = df.merge(library_df, left_on=lookup_col, right_on="EUR Item No.", how="left")
            merged = merged[[quantity_col, "Product", "EUR Item No.", "GBP Item No.", "APMEA Item No.", "USD Pattern No."]]
            merged.to_excel(writer, index=False)
    return buffer.getvalue()

def generate_masterdata(user_data, lookup_col, library_df, master_df):
    """ Generate an Excel file with master data """
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        for sheet, df in user_data.items():
            df.columns = df.columns.str.strip().str.lower()
            quantity_col = next((col for col in df.columns if "quantity" in col.lower() or "qty" in col.lower()), None)
            merged = df.merge(library_df, left_on=lookup_col, right_on="EUR Item No.", how="left")
            merged.to_excel(writer, sheet_name="Detailed Product List", index=False)
            
            master_merged = df.merge(master_df, left_on=lookup_col, right_on="ITEM NO.", how="left")
            master_merged.to_excel(writer, sheet_name="Master Data", index=False)
    return buffer.getvalue()
