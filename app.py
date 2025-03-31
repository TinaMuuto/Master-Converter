import streamlit as st
import pandas as pd
import openpyxl
import os
import re
from io import BytesIO
from docx import Document

#####################
# 1. Data-load functions
#####################

def load_library_data(library_path="Library_data.xlsx"):
    """
    Loads Library_data.xlsx, trims and converts column names to uppercase.
    Expected headers (in row 1) include:
        - Product
        - EUR item no.
        - GBP item no.
        - APMEA item no.
        - USD pattern no.
        - Match Status
    """
    if not os.path.exists(library_path):
        st.error(f"Filen {library_path} mangler i mappen. Upload eller placér filen korrekt.")
        return None
    try:
        df = pd.read_excel(library_path, engine="openpyxl")
        df.columns = df.columns.str.strip().str.upper()
        if "EUR ITEM NO." in df.columns:
            df["EUR ITEM NO."] = df["EUR ITEM NO."].astype(str).str.strip().str.upper()
        return df
    except Exception as e:
        st.error(f"Fejl ved indlæsning af {library_path}: {e}")
        return None

def load_master_data(master_path="Muuto_Master_Data_CON_January_2025_EUR.xlsx"):
    """
    Loads the entire master data file.
    Expected unique lookup column: ITEM NO. (in column B).
    Returns a DataFrame with all columns.
    """
    if not os.path.exists(master_path):
        st.error(f"Filen {master_path} mangler i mappen. Upload eller placér filen korrekt.")
        return None
    try:
        df = pd.read_excel(master_path, engine="openpyxl")
        df.columns = df.columns.str.strip().str.upper()
        if "ITEM NO." in df.columns:
            df["ITEM NO."] = df["ITEM NO."].astype(str).str.strip().str.upper()
        return df
    except Exception as e:
        st.error(f"Fejl ved indlæsning af {master_path}: {e}")
        return None

def load_user_file(uploaded_file):
    """
    Loads the user's uploaded file (pCon-export).
    - If Excel: looks for sheet "Article List", skips the first 2 rows, no header.
    - If CSV: first tries sep=';', if that fails, uses sep=','.
      In both cases, skiprows=2, header=None.
    Returns a DataFrame (without column names) or None.
    """
    try:
        file_name = uploaded_file.name.lower()
        if file_name.endswith(".csv"):
            try:
                df = pd.read_csv(uploaded_file, sep=';', engine="python", header=None, skiprows=2)
            except:
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file, sep=',', engine="python", header=None, skiprows=2)
            return df
        else:
            excel_obj = pd.ExcelFile(uploaded_file, engine="openpyxl")
            if "Article List" in excel_obj.sheet_names:
                df = pd.read_excel(excel_obj, sheet_name="Article List", skiprows=2, header=None)
                return df
            else:
                st.error("Filen indeholder ikke en fane ved navn 'Article List'.")
                return None
    except Exception as e:
        st.error(f"Fejl ved læsning af fil: {e}")
        return None

#####################
# 2. Preprocessing of user data
#####################

def preprocess_user_data(df):
    """
    Extracts columns (by index):
      - Index 17 -> ARTICLE_NO
      - Index 30 -> QUANTITY
      - Index 2  -> SHORT_TEXT
      - Index 4  -> VARIANT_TEXT
    Assumes the first 2 rows have been skipped.
    Replaces NaN in VARIANT_TEXT with an empty string.
    """
    if df.shape[1] < 31:
        st.error("Den uploadede fil indeholder ikke nok kolonner (mindst 31 kræves). Tjek format.")
        return None

    article_no = df.iloc[:, 17].astype(str).str.strip().str.upper()
    quantity = df.iloc[:, 30]
    short_text = df.iloc[:, 2].astype(str).str.strip().str.upper()
    variant_text = df.iloc[:, 4].fillna("").astype(str).str.strip().str.upper()
    
    out_df = pd.DataFrame({
        "ARTICLE_NO": article_no,
        "QUANTITY": quantity,
        "SHORT_TEXT": short_text,
        "VARIANT_TEXT": variant_text
    })
    out_df = out_df[out_df["ARTICLE_NO"].astype(bool)]
    return out_df

#####################
# Helper: Fallback key generation
#####################

def get_fallback_key(article):
    """
    Returns a fallback key for an article number.
    First, splits the article on '-' and takes the first segment.
    Then, if that segment starts with "SPECIAL", removes the "SPECIAL" prefix and any leading spaces.
    Returns the cleaned key in uppercase.
    """
    article = article.strip()
    key = article.split('-')[0].strip().upper()
    if key.startswith("SPECIAL"):
        key = key[len("SPECIAL"):].strip().upper()
    return key

#####################
# 3. Product list for presentations (Word) - using fallback logic
#####################

def generate_presentation_word(df_user, df_library):
    """
    For each row in df_user:
      - Attempts a direct match between ARTICLE_NO and df_library['EUR ITEM NO.'].
      - If a direct match is found, outputs "QUANTITY X PRODUCT".
      - If no direct match is found, computes a fallback key using get_fallback_key and tries to find a match.
      - If a fallback match is found, outputs "QUANTITY X PRODUCT" using the fallback result.
      - Otherwise, outputs "QUANTITY X SHORT_TEXT - VARIANT_TEXT" 
        (omitting '- VARIANT_TEXT' if empty or equals "LIGHT OPTION: OFF").
      Additionally, if the key used for matching contains "ALL COLORS", the match is ignored.
    The list is sorted alphabetically (case-insensitive) before generating a Word document.
    """
    required_cols = ["PRODUCT", "EUR ITEM NO."]
    for col in required_cols:
        if col not in df_library.columns:
            st.error(f"Library_data mangler kolonnen '{col}'. Kan ikke generere præsentationsliste.")
            return None

    lookup_library = df_library.set_index("EUR ITEM NO.")["PRODUCT"].to_dict()
    lines_info = []
    for _, row in df_user.iterrows():
        article_no = row["ARTICLE_NO"]
        quantity = row["QUANTITY"]
        short_text = row["SHORT_TEXT"]
        variant_text = row["VARIANT_TEXT"]
        
        # Forsøg direkte match
        product_match = lookup_library.get(article_no)
        key_used = article_no  # Gemmer den nøgle, der blev brugt til opslaget
        
        # Hvis intet direkte match, anvend fallback-nøgle
        if not product_match:
            fallback_key = get_fallback_key(article_no)
            product_match = lookup_library.get(fallback_key)
            key_used = fallback_key
        
        # Hvis den benyttede nøgle indeholder "ALL COLORS", ignorer matchet
        if product_match and "ALL COLORS" in key_used.upper():
            product_match = None
        
        if product_match:
            sort_key = product_match
            final_line = f"{quantity} X {product_match}"
        else:
            sort_key = short_text
            if variant_text and variant_text != "LIGHT OPTION: OFF":
                final_line = f"{quantity} X {short_text} - {variant_text}"
            else:
                final_line = f"{quantity} X {short_text}"
        lines_info.append((sort_key.upper(), final_line.upper()))
    lines_info.sort(key=lambda x: x[0])
    buffer = BytesIO()
    doc = Document()
    doc.add_heading('Product List for Presentations', level=1)
    for _, line_text in lines_info:
        doc.add_paragraph(line_text)
    doc.save(buffer)
    buffer.seek(0)
    return buffer


#####################
# 4. Order import file (Excel with 2 columns, no header) - using fallback for ARTICLE_NO
#####################

def generate_order_import_excel(df_user):
    """
    Returns an Excel file (as BytesIO) with 2 columns (no headers):
      - Column A: QUANTITY
      - Column B: ARTICLE_NO (cleaned using fallback logic)
    """
    df_order = df_user.copy()
    df_order["ARTICLE_NO"] = df_order["ARTICLE_NO"].apply(get_fallback_key)
    buffer = BytesIO()
    temp_df = df_order[["QUANTITY", "ARTICLE_NO"]].copy()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        temp_df.to_excel(writer, index=False, header=False)
    buffer.seek(0)
    return buffer

#####################
# 5. SKU mapping & Masterdata (with fallback and special-case handling)
#####################

def generate_sku_masterdata_excel(df_user, df_library, df_master):
    """
    Generates an Excel file with two sheets:
    
    1) "Item number mapping":
       - Attempts a direct match between df_user's ARTICLE_NO and Library_data's EUR ITEM NO.
       - If no direct match is found, computes a fallback key using get_fallback_key and attempts a match.
       - Returns the following columns:
           • Quantity in setting (from df_user's QUANTITY)
           • Article No.
           • Short Text
           • Variant text (with NaN replaced by an empty string)
           • Product in setting (from Library_data's Product)
           • EUR item no.
           • GBP item no.
           • APMEA item no.
           • USD pattern no.
           • Match status
    
    2) "Master data export":
       - Attempts a direct match between df_user's ARTICLE_NO and masterdata's ITEM NO.
       - If no direct match is found, computes a fallback key using get_fallback_key and attempts a match.
       - Returns all columns from the masterdata file plus the df_user columns:
         Article No., Short Text, and Variant text (with Variant text cleaned of NaN values).
    """
    # --- ITEM NUMBER MAPPING ---
    rename_map = {
        "PRODUCT": "LIB_PRODUCT",
        "EUR ITEM NO.": "LIB_EUR_ITEM_NO",
        "GBP ITEM NO.": "LIB_GBP_ITEM_NO",
        "APMEA ITEM NO.": "LIB_APMEA_ITEM_NO",
        "USD PATTERN NO.": "LIB_USD_PATTERN_NO",
        "MATCH STATUS": "LIB_MATCH_STATUS"
    }
    df_library_renamed = df_library.rename(columns=rename_map, errors="ignore")
    
    if "LIB_EUR_ITEM_NO" in df_library_renamed.columns:
        merged_direct = pd.merge(
            df_user,
            df_library_renamed,
            how="left",
            left_on="ARTICLE_NO",
            right_on="LIB_EUR_ITEM_NO"
        )
    else:
        merged_direct = df_user.copy()
    
    merged_direct["VARIANT_TEXT"] = merged_direct["VARIANT_TEXT"].fillna("")
    
    df_user_fallback = df_user.copy()
    df_user_fallback["BASE_ARTICLE"] = df_user_fallback["ARTICLE_NO"].apply(get_fallback_key)
    
    if "LIB_EUR_ITEM_NO" in df_library_renamed.columns:
        fallback_merge = pd.merge(
            df_user_fallback,
            df_library_renamed,
            how="left",
            left_on="BASE_ARTICLE",
            right_on="LIB_EUR_ITEM_NO"
        )
    else:
        fallback_merge = df_user_fallback.copy()
    
    for col in ["LIB_PRODUCT", "LIB_EUR_ITEM_NO", "LIB_GBP_ITEM_NO", 
                "LIB_APMEA_ITEM_NO", "LIB_USD_PATTERN_NO", "LIB_MATCH_STATUS"]:
        if col in fallback_merge.columns:
            merged_direct[col] = merged_direct[col].combine_first(fallback_merge[col])
    
    item_number_mapping_df = pd.DataFrame({
        "Quantity in setting": merged_direct["QUANTITY"],
        "Article No.": merged_direct["ARTICLE_NO"],
        "Short Text": merged_direct["SHORT_TEXT"],
        "Variant text": merged_direct["VARIANT_TEXT"].fillna(""),
        "Product in setting": merged_direct.get("LIB_PRODUCT", None),
        "EUR item no.": merged_direct.get("LIB_EUR_ITEM_NO", None),
        "GBP item no.": merged_direct.get("LIB_GBP_ITEM_NO", None),
        "APMEA item no.": merged_direct.get("LIB_APMEA_ITEM_NO", None),
        "USD pattern no.": merged_direct.get("LIB_USD_PATTERN_NO", None),
        "Match status": merged_direct.get("LIB_MATCH_STATUS", None)
    })
    item_number_mapping_df = item_number_mapping_df[item_number_mapping_df["Article No."].astype(bool)]
    
    # --- MASTER DATA EXPORT ---
    if "ITEM NO." not in df_master.columns:
        master_data_export_df = pd.DataFrame(columns=["Article No.", "Short Text", "Variant text"] + df_master.columns.tolist())
    else:
        master_direct = pd.merge(
            df_user,
            df_master,
            how="left",
            left_on="ARTICLE_NO",
            right_on="ITEM NO."
        )
        df_user_master = df_user.copy()
        df_user_master["BASE_ARTICLE"] = df_user_master["ARTICLE_NO"].apply(get_fallback_key)
    
        fallback_master = pd.merge(
            df_user_master,
            df_master,
            how="left",
            left_on="BASE_ARTICLE",
            right_on="ITEM NO."
        )
        for col in df_master.columns:
            master_direct[col] = master_direct[col].combine_first(fallback_master[col])
        master_direct.drop_duplicates(inplace=True)
        master_direct.rename(columns={
            "ARTICLE_NO": "Article No.",
            "SHORT_TEXT": "Short Text",
            "VARIANT_TEXT": "Variant text"
        }, inplace=True)
        master_direct["Variant text"] = master_direct["Variant text"].fillna("")
        front_cols = ["Article No.", "Short Text", "Variant text"]
        other_cols = [c for c in master_direct.columns if c not in front_cols]
        master_data_export_df = master_direct[front_cols + other_cols]
        master_data_export_df = master_data_export_df[master_data_export_df["Article No."].astype(bool)]
    
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        item_number_mapping_df.to_excel(writer, sheet_name="Item number mapping", index=False)
        master_data_export_df.to_excel(writer, sheet_name="Master data export", index=False)
    buffer.seek(0)
    return buffer

#####################
# 6. Streamlit-app
#####################

def main():
    st.set_page_config(page_title="Muuto Product List Generator", layout="centered")
    st.title('Muuto Product List Generator')
    st.write("""
    This tool is designed to **help you structure, validate, and enrich pCon product data effortlessly**.
    
    ### How it works:
    1. Export your product list from pCon (formatted like the example file).
    2. Upload your pCon file to the app.
    3. Click one of the three buttons to generate the file you need.
    4. Once generated, a new button will appear for you to download the file.
    
    ### Expected Outputs:
    - **Product list for presentations:** A Word file with product quantities and descriptions for easy copy-pasting into PowerPoint.
      - File name: product-list_presentation.docx
    - **Order import file:** An Excel file with two columns (Quantity and Article No.) for direct import into the partner platform.
      - File name: order-import.xlsx
    - **Product SKU mapping:** An Excel file with two sheets:
      1. **Product SKU mapping:** Combines uploaded data (Article No. from the "Article List" sheet) with Library_data (matched on EUR item no.). If a direct match is not found, the article number is cleaned using fallback logic (split on '-' and removal of "SPECIAL" prefix) and then matched.
      2. **Master data export:** Uses uploaded data (Article No.) to find matching records in the master data file (matched on ITEM NO.). If no direct match is found, the article number is cleaned similarly and then matched. All master data columns are returned, along with the uploaded file’s Article No., Short Text, and Variant Text.
      - File name: SKUmapping-masterdata.xlsx

    [Download an example file](https://raw.githubusercontent.com/TinaMuuto/Master-Converter/f280308cf9991b7eecb63e44ecac52dfb49482cf/pCon%20-%20exceleksport.xlsx)
    """)
    
    df_library = load_library_data()
    df_master = load_master_data()
    if (df_library is None) or (df_master is None):
        return
    
    uploaded_file = st.file_uploader("Upload your product list (Excel or CSV)", type=['xlsx', 'xls', 'csv'])
    if uploaded_file:
        df_user_raw = load_user_file(uploaded_file)
        if df_user_raw is not None:
            df_user = preprocess_user_data(df_user_raw)
            if df_user is None:
                return
            if st.button("Generate List for presentations"):
                word_buffer = generate_presentation_word(df_user, df_library)
                if word_buffer:
                    st.download_button("Download Word file", data=word_buffer, file_name="product-list.docx")
            if st.button("Generate product list for order import in partner platform"):
                order_buffer = generate_order_import_excel(df_user)
                st.download_button("Download Excel file", data=order_buffer, file_name="order-import.xlsx")
            if st.button("Generate SKU mapping & masterdata"):
                sku_buffer = generate_sku_masterdata_excel(df_user, df_library, df_master)
                if sku_buffer:
                    st.download_button("Download Excel file", data=sku_buffer, file_name="SKUmapping-masterdata.xlsx")

if __name__ == "__main__":
    main()
